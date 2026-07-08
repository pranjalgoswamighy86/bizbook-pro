#!/usr/bin/env python3
"""
BizBook Pro — Backend Team Reference Document Generator
Generates a comprehensive, classified Word document covering the full
development history from v0.0.0 to v6.18.0 / desktop v2.3.0.

CLASSIFICATION: INTERNAL USE ONLY — BACKEND TEAM
Contains: version history, deployment URLs, env var names, security
incidents, architecture decisions, known issues, API structure.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = "/home/z/my-project/download/BizBook_Pro_Backend_Reference.docx"

doc = Document()

# ---- Page setup ----
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# ---- Default font ----
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
style.paragraph_format.line_spacing = 1.3
style.paragraph_format.space_after = Pt(6)

# ---- Helpers ----
def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def add_heading(text, level=1, color=(16, 110, 88)):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*color)
        run.font.name = 'Calibri'
    return h

def add_para(text, bold=False, italic=False, color=None, size=11, indent=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.8 + level * 0.6)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(30, 41, 59)
    # light gray background
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'F1F5F9')
    pPr.append(shd)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_bg(cell, '0F766E')
    # rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '10B981')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_page_break():
    doc.add_page_break()

# ============================================================
# COVER PAGE
# ============================================================
add_para("CLASSIFIED — INTERNAL USE ONLY", bold=True, color=(220, 38, 38), size=14)
add_para("BACKEND TEAM REFERENCE DOCUMENT", bold=True, color=(16, 110, 88), size=12)
add_hr()
doc.add_paragraph()
add_para("BizBook Pro", bold=True, size=36, color=(15, 23, 42))
add_para("Complete Development Reference", bold=True, size=20, color=(16, 110, 88))
add_para("Version 0.0.0 → v6.18.0 (Web) / v2.3.0 (Desktop)", size=14, color=(71, 85, 105))
doc.add_paragraph()
add_para("A Product by Tahigo International", size=13, italic=True, color=(100, 116, 139))
add_para("Guwahati, Assam, India  ·  www.tahigo.in", size=12, color=(100, 116, 139))
doc.add_paragraph()
add_hr()
add_para("DOCUMENT METADATA", bold=True, size=11, color=(16, 110, 88))
add_para("Document Version: 1.0", size=10)
add_para("Generated: 2026-07-08", size=10)
add_para("Classification: INTERNAL — BACKEND TEAM ONLY", size=10, bold=True, color=(220, 38, 38))
add_para("Distribution: Restricted to Tahigo International backend engineering team", size=10)
add_para("Source: Complete development chat log (complete chat 05-07-2026 11.30 pm.txt, 66,865 lines)", size=10)
doc.add_paragraph()
add_hr()
add_para("WARNING: This document contains deployment URLs, environment variable names, security incident details, architecture decisions, and known issues. Do NOT share outside the backend team. Do NOT commit to public repositories. Rotate any credentials that may have been exposed in git history.", bold=True, size=10, color=(220, 38, 38))
add_page_break()

# ============================================================
# TABLE OF CONTENTS (manual)
# ============================================================
add_heading("Table of Contents", level=1)
toc_items = [
    "1. Executive Summary",
    "2. Product Overview & Architecture",
    "3. Technology Stack",
    "4. Deployment Infrastructure",
    "5. Environment Variables Reference",
    "6. Version History (v0.0.0 → v6.18.0)",
    "7. Desktop Application History (v2.0.0 → v2.3.0)",
    "8. Database Schema & Models",
    "9. API Endpoints Catalog",
    "10. Security Architecture & Incidents",
    "11. Role-Based Access Control",
    "12. Subscription & Billing Logic",
    "13. Known Issues & Technical Debt",
    "14. Critical Files Reference",
    "15. Operational Runbook",
    "16. Security Rotation Checklist",
]
for item in toc_items:
    add_para(item, size=11)
add_page_break()

# ============================================================
# 1. EXECUTIVE SUMMARY
# ============================================================
add_heading("1. Executive Summary", level=1)
add_para("BizBook Pro is a multi-tenant SaaS billing, inventory, accounting, and payroll platform built for Indian small and medium enterprises. The product is developed by Tahigo International (Guwahati, Assam, India) and is currently live in production at version v6.18.0 (web) and v2.3.0 (Electron desktop).")
add_para("This document is the single source of truth for the backend engineering team. It consolidates every architectural decision, deployment detail, security incident, version milestone, and operational procedure gathered across the full development lifecycle — from the initial request on 2026-06-15 through to the latest v6.18.0 deploy on 2026-07-08.")
add_para("Key facts at a glance:", bold=True)
add_bullet("Web app: v6.18.0 — live on Railway at https://carefree-success-production-7766.up.railway.app/")
add_bullet("Desktop app: v2.3.0 — Electron 33, distributed via GitHub Releases at https://github.com/pranjalgoswamighy86/bizbook-pro")
add_bullet("Codebase: 29+ business modules, 42 REST API endpoints, 30+ Prisma database models")
add_bullet("Performance: 926 requests/sec with 0 failures under stress test")
add_bullet("Stack: Next.js 16, React 19, TypeScript, Tailwind CSS 4, Prisma ORM 6, PostgreSQL, Electron 33")
add_bullet("Active tenants: 11 (as of 2026-07-08 startup log)")
add_bullet("Total commits: 200+ across 23 days of active development")
add_page_break()

# ============================================================
# 2. PRODUCT OVERVIEW & ARCHITECTURE
# ============================================================
add_heading("2. Product Overview & Architecture", level=1)

add_heading("2.1 Origin & Vision", level=2)
add_para("BizBook Pro was conceived on 2026-06-15 in response to the founder's frustration with the complexity of existing Indian accounting software (Tally Prime, Marg ERP). The core thesis: build a cloud-native, mobile-responsive, AI-powered ERP that is dramatically easier to use than legacy desktop tools, while remaining GST-compliant and affordable via hour-based subscription pricing.")

add_heading("2.2 High-Level Architecture", level=2)
add_para("The system is a single-page Next.js application with API routes served from the same codebase. The same codebase is packaged as an Electron desktop app — there is one codebase, two delivery surfaces (web + desktop).")
add_bullet("Frontend: Next.js 16 (App Router), React 19, TypeScript strict mode, Tailwind CSS 4, shadcn/ui components, Zustand for state management, TanStack Query for server state.")
add_bullet("Backend: Next.js API Routes (42 endpoints), Prisma ORM 6, PostgreSQL 14+ on Railway.")
add_bullet("Desktop: Electron 33 wraps the Railway URL directly (production mode loads https://carefree-success-production-7766.up.railway.app/ — no local Next.js server in the desktop app).")
add_bullet("AI Layer: 4-provider fallback chain — Gemini → OpenAI → DeepSeek → Anthropic → ZAI (in-house). All AI processing is server-side.")
add_bullet("Auth: JWT session cookies (bizbook_session) + Email/SMS OTP for registration and password reset. 5-tier role-based access control (4 tenant roles + 1 backend-only platform role).")
add_bullet("Deployment: Docker container on Railway, PM2 cluster mode (2 instances) for zero-downtime. GitHub Actions builds the desktop installer for Windows, macOS, and Linux.")

add_heading("2.3 Multi-Tenancy Model", level=2)
add_para("Every database table carries a tenantId column. All API routes enforce tenant isolation via requireAuthAndTenant() or requireAuthAndRole() middleware — the access.tenantId is sourced from the authenticated JWT, never from the request body. A tenant protection safeguard runs on every startup, querying all registered tenants and asserting they are all marked ACTIVE (none soft-deleted).")
add_para("As of 2026-07-08, 11 active tenants are registered:")
add_code_block(
    "admin@bizbook.pro (BizBook Pro, free)\n"
    "GOSWAMIPRANJALGHY86@GMAIL.COM (TEST, free)\n"
    "bakersmartghy@gmail.com (BAKERS MART, free)\n"
    "kdhomesghy@gmail.com (Test, free)\n"
    "pritisharma6364@gmail.com (Na, free)\n"
    "amritsonowal165@gmail.com (ZTX P LTD, free)\n"
    "thewhiskeyghy@gmail.com (The Whiskey, free)\n"
    "bakersmartghy@gmail.com (Bakers Mart - DMP, free)\n"
    "pranjalgoswamighy86@gmail.com (Tahigo International, free)\n"
    "bd10052024@outlook.com (test, free)\n"
    "sarmahirak389@gmail.com (Sarma store, free)"
)
add_para("NOTE: The above tenant list is operational data, not test data. All 11 tenants are real registered businesses. The tenant-protection safeguard prevents accidental deletion of any of these.", italic=True, color=(220, 38, 38))
add_page_break()

# ============================================================
# 3. TECHNOLOGY STACK
# ============================================================
add_heading("3. Technology Stack", level=1)

add_heading("3.1 Frontend", level=2)
add_table(
    ["Layer", "Technology", "Version", "Purpose"],
    [
        ["Framework", "Next.js", "16.1.3", "App Router, API Routes, SSR/SSG"],
        ["UI Library", "React", "19.0.0", "Component rendering"],
        ["Language", "TypeScript", "5.x", "Type safety (strict mode)"],
        ["Styling", "Tailwind CSS", "4.x", "Utility-first CSS"],
        ["Components", "shadcn/ui + Radix UI", "latest", "Accessible primitives"],
        ["State", "Zustand", "5.0.6", "Client state (persisted)"],
        ["Server State", "TanStack Query", "5.82.0", "Data fetching cache"],
        ["Forms", "React Hook Form + Zod", "7.60 / 4.0", "Validation"],
        ["Charts", "Recharts", "2.15.4", "Dashboard visualizations"],
        ["Tables", "TanStack Table", "8.21.3", "Sortable, filterable grids"],
        ["Icons", "lucide-react", "0.525.0", "UI icons"],
        ["Date", "date-fns", "4.1.0", "Date manipulation"],
        ["Animations", "framer-motion", "12.23.2", "Transitions"],
    ],
    col_widths=[3, 4, 2.5, 6]
)

add_heading("3.2 Backend", level=2)
add_table(
    ["Layer", "Technology", "Version", "Purpose"],
    [
        ["Runtime", "Node.js", "20-slim (Docker)", "Server runtime"],
        ["Framework", "Next.js API Routes", "16.1.3", "REST API (42 endpoints)"],
        ["ORM", "Prisma", "6.19.3", "Database access + migrations"],
        ["Database", "PostgreSQL", "14+ (Railway)", "Primary datastore"],
        ["Process Mgr", "PM2", "5.4.3", "Cluster mode (2 instances)"],
        ["Auth", "JWT (cookie) + OTP", "—", "Session + 2FA"],
        ["Validation", "Zod", "4.0.2", "Schema validation"],
        ["Payments", "Razorpay", "2.9.6", "Subscription checkout"],
        ["Email", "Brevo + Resend + Nodemailer", "latest", "OTP + notifications"],
        ["SMS", "2Factor.in / custom", "—", "OTP fallback"],
        ["Excel", "SheetJS (xlsx)", "0.18.5", "Backup export/import"],
        ["PDF", "pdf-parse + mammoth", "2.4.5 / 1.12.0", "Invoice parsing"],
    ],
    col_widths=[3, 4, 2.5, 6]
)

add_heading("3.3 Desktop (Electron)", level=2)
add_table(
    ["Layer", "Technology", "Version", "Purpose"],
    [
        ["Shell", "Electron", "33.x", "Desktop wrapper"],
        ["Builder", "electron-builder", "25.1.8", "NSIS / DMG / AppImage"],
        ["Packaging", "asar", "enabled", "Source protection"],
        ["Auto-update", "electron-updater", "via GitHub", "Release-based updates"],
        ["Fingerprint", "SecuGen / DigitalPersona SDK", "—", "Biometric attendance"],
        ["Printing", "ESC/POS raw + webContents.print", "—", "Thermal + A4 silent print"],
    ],
    col_widths=[3, 4, 2.5, 6]
)

add_heading("3.4 AI Layer — 4-Provider Fallback", level=2)
add_para("The AI layer (src/lib/multi-ai.ts) tries providers in this order. If a provider fails (network, auth, rate limit), it falls back to the next:")
add_bullet("1. Google Gemini (primary) — GEMINI_API_KEY env var")
add_bullet("2. OpenAI — OPENAI_API_KEY env var")
add_bullet("3. DeepSeek — DEEPSEEK_API_KEY env var")
add_bullet("4. Anthropic — ANTHROPIC_API_KEY env var")
add_bullet("5. ZAI (in-house fallback) — always available as last resort")
add_para("Used by: AI Smart Import (photo/Excel to entry), AI Business Valuation, AI Support Chat, Invoice Analysis, Smart Search. All AI processing is server-side; no business data is retained by providers — only transient prompts.")
add_page_break()

# ============================================================
# 4. DEPLOYMENT INFRASTRUCTURE
# ============================================================
add_heading("4. Deployment Infrastructure", level=1)

add_heading("4.1 Railway (Web App)", level=2)
add_bullet("Project name: carefree-success")
add_bullet("Live URL: https://carefree-success-production-7766.up.railway.app/")
add_bullet("Internal URL: https://carefree-success-production-7766.up.railway.app (same)")
add_bullet("Database: PostgreSQL on Railway (postgres.railway.internal:5432, database 'railway')")
add_bullet("Volume: /var/lib/containers/railwayapp/bind-mounts/640b712b-8189-4e98-8fd3-b1a7fd2892c2/vol_l1bcuqbizzosgjgf")
add_bullet("Build: Dockerfile (node:20-slim base, 9-step build, npm install + prisma generate + next build + postbuild)")
add_bullet("Start command: scripts/railway-start.js (regenerates Prisma client, syncs schema, runs tenant protection, creates startup backup, starts PM2 cluster)")
add_bullet("PM2 config: ecosystem.config.js (2 instances, no daemon mode, falls back to direct server.js if PM2 fails)")
add_bullet("Port: 8080 (Railway maps to 443 externally)")
add_bullet("Health check: /api/health endpoint")
add_para("NOTE: An earlier Railway project named 'bizbook-pro' exists but failed. The 'carefree-success' project is the active deployment. Do NOT delete either project without confirming which is live.", italic=True, color=(220, 38, 38))

add_heading("4.2 GitHub Repository", level=2)
add_bullet("URL: https://github.com/pranjalgoswamighy86/bizbook-pro")
add_bullet("Visibility: PUBLIC (intentional — enables free GitHub Actions minutes for desktop builds)")
add_bullet("Default branch: main")
add_bullet("Auto-deploy: Railway watches main branch, auto-rebuilds on push (~3-5 min build time)")
add_bullet("Desktop builds: Triggered by git tag push (v* pattern) via .github/workflows/build-desktop.yml")
add_bullet("Release process: git tag v2.x.x && git push origin v2.x.x → GitHub Actions builds Windows/macOS/Linux → attaches installers to GitHub Release (draft)")

add_heading("4.3 Desktop App Distribution", level=2)
add_table(
    ["Platform", "Format", "Architecture", "Distribution"],
    [
        ["Windows", "NSIS installer + Portable .exe", "x64", "GitHub Releases"],
        ["macOS", "DMG + ZIP", "x64 + arm64", "GitHub Releases"],
        ["Linux", "AppImage + DEB", "x64", "GitHub Releases"],
    ],
    col_widths=[3, 5, 3, 4]
)
add_para("Desktop app behavior: In production mode, the Electron app loads the Railway URL directly (no local Next.js server). This was a v5.6 fix — the previous approach of running a local server failed because DATABASE_URL wasn't available in the desktop environment. The desktop app clears its HTTP cache + storage on every launch (v2.3.0 fix) to ensure it always loads the latest Railway deploy.")
add_page_break()

# ============================================================
# 5. ENVIRONMENT VARIABLES REFERENCE
# ============================================================
add_heading("5. Environment Variables Reference", level=1)
add_para("All environment variables are set in Railway → carefree-success → Variables. The .env, .env.local, .env.production files are gitignored. Only .env.example (placeholders only) is committed.", color=(220, 38, 38), bold=True)

add_heading("5.1 Required Variables", level=2)
add_table(
    ["Variable", "Purpose", "Example / Notes"],
    [
        ["DATABASE_URL", "PostgreSQL connection string", "postgresql://postgres:***@postgres.railway.internal:5432/railway"],
        ["SESSION_SECRET", "JWT session signing secret", "Generate with: openssl rand -hex 32"],
        ["ADMIN_EMAIL", "Initial super-admin email", "admin@bizbook.pro"],
        ["RAZORPAY_KEY_ID", "Razorpay payment gateway key ID", "rzp_live_*** or rzp_test_***"],
        ["RAZORPAY_KEY_SECRET", "Razorpay payment gateway secret", "(rotate — was in old git history)"],
        ["BREVO_API_KEY", "Brevo (Sendinblue) email API key", "xkeysib-*** (rotate — was in old git history)"],
        ["BREVO_FROM_EMAIL", "Sender email for Brevo", "pranjalgoswamighy86@gmail.com"],
        ["BREVO_FROM_NAME", "Sender display name", "BizBook Pro"],
        ["RESEND_API_KEY", "Resend email API (fallback)", "re_***"],
        ["RESEND_FROM", "Resend sender", "BizBook Pro <onboarding@resend.dev>"],
        ["GEMINI_API_KEY", "Google Gemini AI (primary)", "AIza*** (35 chars)"],
        ["OPENAI_API_KEY", "OpenAI (fallback 2)", "sk-***"],
        ["DEEPSEEK_API_KEY", "DeepSeek (fallback 3)", "sk-***"],
        ["ANTHROPIC_API_KEY", "Anthropic Claude (fallback 4)", "sk-ant-***"],
        ["SMTP_HOST", "SMTP fallback host", "smtp.gmail.com"],
        ["SMTP_PORT", "SMTP fallback port", "465"],
        ["SMTP_USER", "SMTP username", "pranjalgoswamighy86@gmail.com"],
        ["SMTP_PASS", "SMTP password (Gmail App Password)", "16-char app password"],
    ],
    col_widths=[4, 5, 6]
)

add_heading("5.2 Optional / Feature-Specific Variables", level=2)
add_table(
    ["Variable", "Purpose"],
    [
        ["SMS_API_KEY", "2Factor.in SMS API key (for OTP via SMS)"],
        ["WHATSAPP_VERSION", "WhatsApp Cloud API version (default v18.0)"],
        ["UPCHECKOUT_API_KEY", "UPI checkout verification (if used)"],
        ["NODE_OPTIONS", "--max-old-space-size=2048 (build) / 512 (runtime OOM fix)"],
        ["PORT", "8080 (Railway) / 3000 (dev) / 3456 (desktop prod fallback)"],
    ],
    col_widths=[4, 11]
)

add_heading("5.3 Security Note — Leaked Keys in Git History", level=2, color=(220, 38, 38))
add_para("CRITICAL: The following keys were committed to git history (public repo) and MUST be rotated:", bold=True, color=(220, 38, 38))
add_bullet("BREVO_API_KEY — leaked in commit a094565 (2026-06-20) in a documentation comment. Value: [REDACTED — full Brevo API key is in Git history, must be rotated at Brevo dashboard]")
add_bullet("RAZORPAY_KEY_SECRET — may have been in early commits (verify with git log -p | grep -i razorpay)")
add_para("Rotation steps:", bold=True)
add_bullet("Brevo: Visit https://app.brevo.com/settings/keys/api — delete old key, create new, update Railway env var")
add_bullet("Razorpay: Visit https://dashboard.razorpay.com/app/keys — regenerate key, update Railway env var, update any hardcoded references")
add_bullet("After rotation: Force re-deploy on Railway (Settings → Redeploy) so the new env vars take effect")
add_para("NOTE: Git history cannot be rewritten without force-pushing and coordinating with all contributors. The keys are already public. Rotation is the only remediation.", italic=True, color=(220, 38, 38))
add_page_break()

# ============================================================
# 6. VERSION HISTORY
# ============================================================
add_heading("6. Version History (v0.0.0 → v6.18.0)", level=1)
add_para("The following table consolidates every major version milestone extracted from the development chat log. Dates are approximate based on commit timestamps.")

add_heading("6.1 Phase 1 — Foundation (v0.0.0 → v1.0.0)", level=2)
add_table(
    ["Version", "Date", "Milestone"],
    [
        ["v0.0.0", "2026-06-15", "Initial request: multi-tenant billing + inventory SaaS for Indian SMEs"],
        ["v0.1.0", "2026-06-15", "Prisma schema, API routes, sidebar, dashboard, sale/purchase/expense registers"],
        ["v0.2.0", "2026-06-15", "Inventory, bank statement, P&L, day report, balance sheet modules"],
        ["v0.3.0", "2026-06-15", "Staff & salary, debtors, creditors, payments, receipts, settings"],
        ["v0.4.0", "2026-06-15", "Subscription/tenant management, Excel export, periodic search filters"],
        ["v1.0.0", "2026-06-16", "package.json renamed to bizbook-pro-server v1.0.0, deploy scripts added"],
    ],
    col_widths=[2, 2.5, 10.5]
)

add_heading("6.2 Phase 2 — Polish & Deployment (v4.0 → v4.60)", level=2)
add_table(
    ["Version", "Date", "Milestone"],
    [
        ["v4.43", "2026-06-20", "isDesktop crash fix, email-only OTP, UPI payment verify, Brevo email pipeline"],
        ["v4.49", "2026-06-22", "Help & Support modal, F1 shortcut"],
        ["v4.50", "2026-06-22", "Version display v4.50.0, 'A Product by Tahigo International'"],
        ["v4.52", "2026-06-22", "Mobile-compatible help modal, responsive sizing"],
        ["v4.55.1", "2026-06-22", "Build failure fixed, all optimizations deployed"],
        ["v4.56", "2026-06-23", "PostgreSQL + PM2 cluster startup (startup banner 'v4.56 — PostgreSQL + PM2')"],
        ["v4.56.1", "2026-06-23", "Build failure fixed"],
        ["v4.56.2", "2026-06-23", "4 critical fixes + screenshot analysis"],
        ["v4.57.1", "2026-06-23", "Inventory item type dropdown fixed, settings updated"],
        ["v4.57.2", "2026-06-23", "'Unknown Language' issue fixed (AI prompts forced English)"],
        ["v4.59.1", "2026-06-23", "Dashboard blank cards fixed"],
        ["v4.62.1", "2026-06-24", "Cash customer Balance Due = 0 enforcement"],
        ["v4.64.1", "2026-06-24", "Sidebar visibility fixed, Help label renamed"],
        ["v4.64.2", "2026-06-24", "Desktop sidebar branding fixed"],
        ["v4.64.3", "2026-06-24", "Personal info removed from all user-facing text"],
        ["v4.64.4", "2026-06-24", "Build failure fixed (apostrophe escape)"],
        ["v4.66.1", "2026-06-24", "CardDescription import fix, partPaymentAmount cleanup"],
    ],
    col_widths=[2, 2.5, 10.5]
)

add_heading("6.3 Phase 3 — Feature Expansion (v4.72 → v4.160)", level=2)
add_table(
    ["Version", "Date", "Milestone"],
    [
        ["v4.72", "2026-06-25", "Payment mode breakdown: CASH, UPI, CARD, PART_PAYMENT, OTHERS"],
        ["v4.106", "2026-06-24", "Quotation-to-Confirm workflow, UPI QR shows only UPI amount, upiAmount field"],
        ["v4.107", "2026-06-24", "HelpSupportManagement inline lazy-load"],
        ["v4.113.1", "2026-06-25", "triggerAutoBackup improvements"],
        ["v4.117", "2026-06-25", "Deployed code revert issue identified"],
        ["v4.123", "2026-06-25", "Re-pushed after revert; Settings shows v4.123.0"],
        ["v4.125", "2026-06-25", "Anti-negative value validation (GST compliance)"],
        ["v4.143", "2026-06-26", "discountAmount IS the price customer pays (e.g., ₹150 for 50Hrs)"],
        ["v4.155", "2026-06-26", "Auto Excel backup after every transaction, offline mode banner"],
        ["v4.159", "2026-06-27", "BOM-aware inventory reversal on edit/delete, salary in P&L"],
        ["v4.160", "2026-06-27", "Cash customer rule + subscription 15% surcharge"],
        ["v4.168", "2026-06-29", "Service Worker cache v4.168.0 (forces browser refetch)"],
    ],
    col_widths=[2, 2.5, 10.5]
)

add_heading("6.4 Phase 4 — Invoice & Print Polish (v4.178 → v4.190)", level=2)
add_table(
    ["Version", "Date", "Milestone"],
    [
        ["v4.178", "2026-07-03", "SW cache bumped to force browser reload (was stuck at v4.155)"],
        ["v4.180", "2026-07-03", "Invoice CSS rewrite, deeper/bolder fonts (Courier New, weight 900)"],
        ["v4.187", "2026-07-03", "SW cache bump to invalidate cached preview HTML"],
        ["v4.188", "2026-07-03", "SW cache bump for iframe HTML"],
        ["v4.189", "2026-07-03", "SW cache bump for iframe HTML"],
        ["v4.190", "2026-07-04", "Explicit paper size selection (A4 / Thermal 80mm)"],
    ],
    col_widths=[2, 2.5, 10.5]
)

add_heading("6.5 Phase 5 — Desktop App & Electron (v5.0 → v5.13)", level=2)
add_table(
    ["Version", "Date", "Milestone"],
    [
        ["v5.0.1", "2026-07-04", "Paper-agnostic invoice layout (90% width, adaptive)"],
        ["v5.2.1", "2026-07-04", "Modal positioning fix"],
        ["v5.2.5", "2026-07-04", "Modal anchored to top (root cause: items-center centering)"],
        ["v5.4", "2026-07-05", "Window spawn kill switch (3 windows / 10s = force quit)"],
        ["v5.6", "2026-07-05", "Desktop loads Railway URL directly (no local server)"],
        ["v5.8", "2026-07-05", "Printer auto-detection (58mm / 80mm / A4)"],
        ["v5.10", "2026-07-05", "Improved silent print with error handling"],
        ["v5.12", "2026-07-05", "ESC/POS direct printing (raw commands via copy /b)"],
        ["v5.13", "2026-07-05", "ESC/POS without native deps (Windows copy /b, lpr on macOS/Linux)"],
    ],
    col_widths=[2, 2.5, 10.5]
)

add_heading("6.6 Phase 6 — Subscription, Users, Quotation (v6.0 → v6.18)", level=2)
add_table(
    ["Version", "Date", "Milestone"],
    [
        ["v6.0.1", "2026-07-06", "Overflow fix"],
        ["v6.3", "2026-07-06", "Allow about:blank print preview windows in Electron"],
        ["v6.5", "2026-07-06", "Tenant-level wallet (one wallet, all companies share hours)"],
        ["v6.12", "2026-07-06", "Single main admin lock, global user count, 15% surcharge"],
        ["v6.13.1", "2026-07-07", "Guide popup fixed (createPortal for z-index)"],
        ["v6.14", "2026-07-07", "F1/Ctrl+N/Ctrl+1-5 keyboard shortcuts (web + Electron IPC)"],
        ["v6.14.1", "2026-07-07", "Menu shortcuts verified, deployed"],
        ["v6.14.2", "2026-07-07", "Menu shortcuts fixed, version updated"],
        ["v6.14.3", "2026-07-07", "createPortal + keyboard shortcuts both fixed"],
        ["v6.15.0", "2026-07-08", "AI Smart Import fix (FormData file uploads)"],
        ["v6.16.0", "2026-07-08", "Electron menu bar fix — clear cache on launch + global MenuActionBridge + VersionBadge"],
        ["v6.16.1", "2026-07-08", "Security hardening — whitelist Electron menu actions"],
        ["v6.17.0", "2026-07-08", "F1 opens AI Support Chat (was: admin ticket page)"],
        ["v6.18.0", "2026-07-08", "Cash customer sale confirmation — allow if balance due = 0"],
    ],
    col_widths=[2, 2.5, 10.5]
)
add_page_break()

# ============================================================
# 7. DESKTOP APP HISTORY
# ============================================================
add_heading("7. Desktop Application History (v2.0.0 → v2.3.0)", level=1)
add_para("The Electron desktop app is versioned separately from the web app. The desktop shell version lives in package.json. The desktop app loads the Railway web app at runtime, so web app updates take effect on next desktop launch (cache is cleared on every launch since v2.3.0).")

add_table(
    ["Desktop Version", "Date", "Web App at Time", "Milestone"],
    [
        ["v2.0.0", "2026-07-04", "v5.0.x", "Initial Electron release — Windows .exe, macOS .dmg, Linux .AppImage via GitHub Actions"],
        ["v2.0.1", "2026-07-05", "v5.2.x", "Critical fix pushed"],
        ["v2.0.2", "2026-07-05", "v5.2.x", "Critical fix pushed"],
        ["v2.0.3", "2026-07-05", "v5.2.x", "Two issues found and fixed (macOS only build — Windows/Linux failed Prisma)"],
        ["v2.0.4", "2026-07-05", "v5.2.x", "Fresh build tag to retry Windows/Linux"],
        ["v2.0.7", "2026-07-05", "v5.10+", "Silent print fix — auto-detection confirmed working"],
        ["v2.1.2", "2026-07-07", "v6.9.x", "Fix pushed + build triggered"],
        ["v2.1.3", "2026-07-07", "v6.9.x", "Fix pushed"],
        ["v2.1.4", "2026-07-07", "v6.13.x", "Fix pushed — built BEFORE v6.14 menu fix"],
        ["v2.2.0", "2026-07-08", "v6.14.3", "Menu fix + createPortal + keyboard shortcuts"],
        ["v2.2.1", "2026-07-08", "v6.15.x", "Tag verified (does not include v6.15 AI Import fix — pushed after tag)"],
        ["v2.3.0", "2026-07-08", "v6.16.x", "Menu bar root-cause fix: clear cache on launch + robust sendMenuAction retry + VersionBadge + security whitelist"],
    ],
    col_widths=[2.5, 2.5, 2.5, 7.5]
)
add_page_break()

# ============================================================
# 8. DATABASE SCHEMA & MODELS
# ============================================================
add_heading("8. Database Schema & Models", level=1)
add_para("Database: PostgreSQL on Railway (postgres.railway.internal:5432, database 'railway'). ORM: Prisma 6.19.3. Schema file: prisma/schema.prisma. Soft-delete architecture — records are never hard-deleted (isDeleted + deletedAt fields).")

add_heading("8.1 Core Models (30+ total)", level=2)
add_table(
    ["Model", "Purpose", "Key Fields"],
    [
        ["Tenant", "SaaS tenant (one per business)", "id, name, plan, gstNumber, currency, upiId"],
        ["User", "Authenticated user (belongs to tenant)", "id, email, name, role, tenantId, passwordHash"],
        ["Sale", "Sales invoice / quotation", "id, invoiceNumber, date, partyName, items, subtotal, gstAmount, totalAmount, paymentStatus, invoiceStatus (QUOTATION/CONFIRMED), upiAmount, amountReceived, einvoiceIrn, tenantId, createdBy, isDeleted, version"],
        ["Purchase", "Purchase bill", "Same structure as Sale + creditor side effects"],
        ["Expense", "Operating expense", "id, date, amount, category, paymentMode, description, tenantId"],
        ["InventoryItem", "Stock item", "id, name, hsnCode, itemType (RAW_MATERIAL/FINISHED_PRODUCT/SERVICE), currentStock, purchasePrice, salePrice, mrp, gstRate, value, tenantId"],
        ["Product", "Finished product with BOM", "id, name, ingredients[] (relation to InventoryItem)"],
        ["Party", "Customer or supplier", "id, name, type (CUSTOMER/SUPPLIER), currentBalance, gstNumber, tenantId"],
        ["Debtor", "Accounts receivable", "id, name, currentBalance, openingBalance, tenantId"],
        ["Creditor", "Accounts payable", "Same as Debtor"],
        ["Payment", "Money out", "id, date, amount, paymentMode, partyName, tenantId"],
        ["Receipt", "Money in", "id, date, amount, paymentMode, partyName, reference, tenantId"],
        ["BankTransaction", "Bank statement entry", "id, date, description, deposit, withdrawal, balance, bankName, category, tenantId"],
        ["Staff", "Employee record", "id, name, role, salary, fingerprintTemplate, tenantId"],
        ["SalaryPayment", "Payroll disbursement", "id, staffId, amount, paidDate, month, tenantId"],
        ["Attendance", "Daily attendance", "id, staffId, date, status (PRESENT/ABSENT), checkIn, checkOut, tenantId"],
        ["Account", "Chart of accounts", "id, accountCode, name, type (Asset/Liability/Equity/Revenue/Expense), tenantId"],
        ["JournalEntry", "Double-entry header", "id, entryDate, reference, description, sourceType (SALE/PURCHASE/MANUAL), sourceId, isPosted, tenantId, createdBy"],
        ["JournalLine", "Double-entry line", "id, journalEntryId, accountId, debit, credit, description"],
        ["Subscription", "Tenant subscription", "id, tenantId, planHours, remainingSeconds, planName, isFreeTier, status, roleAllocation"],
        ["AuditLog", "Immutable audit trail", "id, tenantId, userId, userName, action, entityType, entityId, changes (JSON), createdAt"],
        ["HelpTicket", "AI Support Chat tickets", "id, tenantId, userEmail, question, answer, status, createdAt"],
        ["Batch", "Inventory batch (expiry tracking)", "id, inventoryItemId, batchNumber, manufactureDate, expiryDate, quantity, tenantId"],
        ["PriceList", "Customer-specific pricing", "id, name, partyId, items[], tenantId"],
        ["UsageLog", "Per-user hour tracking", "id, userId, tenantId, secondsUsed, loggedAt"],
    ],
    col_widths=[3, 5, 7]
)
add_para("Note: This is a representative subset. The full schema has 30+ models. See prisma/schema.prisma for the canonical definition.", italic=True, size=10)
add_page_break()

# ============================================================
# 9. API ENDPOINTS CATALOG
# ============================================================
add_heading("9. API Endpoints Catalog", level=1)
add_para("All 42 API endpoints live under /api/. Every endpoint enforces requireAuthAndTenant() or requireAuthAndRole() at the top. The access.tenantId comes from the JWT, never from the request body. Responses are JSON.")

add_heading("9.1 Core CRUD Endpoints", level=2)
add_table(
    ["Endpoint", "Methods", "Purpose"],
    [
        ["/api/auth", "POST", "Login, register, send-otp, reset-password, change-password, me, switch-company"],
        ["/api/sales", "POST", "create, update, delete, confirm-sale, list (actions in body)"],
        ["/api/purchases", "POST", "create, update, delete, list"],
        ["/api/expenses", "POST", "create, update, delete, list"],
        ["/api/inventory", "POST", "create, update, delete, list, stock-adjust"],
        ["/api/products", "POST", "create, update, delete, list (with BOM)"],
        ["/api/parties", "POST", "create, update, delete, list (customers/suppliers)"],
        ["/api/payments", "POST", "create, update, delete, list"],
        ["/api/receipts", "POST", "create, update, delete, list"],
        ["/api/bank", "POST", "create, update, delete, list, reconcile"],
        ["/api/staff", "POST", "create, update, delete, list, attendance"],
        ["/api/attendance", "POST", "check-in, check-out, report"],
        ["/api/tds", "POST", "TDS register CRUD"],
        ["/api/creditors", "POST", "Creditors list, aging"],
        ["/api/debtors", "POST", "Debtors list, aging"],
        ["/api/accounts", "POST", "Chart of accounts CRUD"],
        ["/api/journal-entries", "POST", "Manual journal entries, list, reverse"],
        ["/api/ledger", "POST", "General ledger, trial balance"],
        ["/api/price-lists", "POST", "Customer price lists CRUD"],
        ["/api/batches", "POST", "Batch tracking CRUD"],
    ],
    col_widths=[4, 2.5, 8.5]
)

add_heading("9.2 Reports & Analytics", level=2)
add_table(
    ["Endpoint", "Methods", "Purpose"],
    [
        ["/api/reports", "POST", "actions: pnl, day-report, balance-sheet, dashboard"],
        ["/api/audit", "POST", "Audit log query (MAIN_ADMIN+)"],
    ],
    col_widths=[4, 2.5, 8.5]
)

add_heading("9.3 Subscription & Billing", level=2)
add_table(
    ["Endpoint", "Methods", "Purpose"],
    [
        ["/api/subscription", "POST", "actions: get-status, recharge, get-plans, create-free-tier"],
        ["/api/razorpay", "POST", "Create order, verify signature (HMAC-SHA256)"],
        ["/api/upi-checkout", "POST", "UPI checkout verification (3-layer: SMS + IMAP + screenshot)"],
        ["/api/payment-proof", "POST", "Upload payment screenshot for manual review"],
        ["/api/payment-proof/review", "POST", "Super-admin reviews payment proofs"],
    ],
    col_widths=[4, 2.5, 8.5]
)

add_heading("9.4 AI Endpoints", level=2)
add_table(
    ["Endpoint", "Methods", "Purpose"],
    [
        ["/api/ai-import", "POST", "FormData file upload → AI parses → returns structured entry"],
        ["/api/ai-valuation", "POST", "Business valuation (15+ multiples)"],
        ["/api/ai-smart-search", "POST", "Natural-language search across data"],
        ["/api/analyze-invoice", "POST", "Parse invoice photo/PDF"],
        ["/api/help-chat", "POST", "AI Support Chat (F1) — auto-escalates to ticket"],
        ["/api/help-tickets", "POST", "Ticket management (super-admin)"],
    ],
    col_widths=[4, 2.5, 8.5]
)

add_heading("9.5 System & Operations", level=2)
add_table(
    ["Endpoint", "Methods", "Purpose"],
    [
        ["/api/health", "GET", "Health check (used by Railway + monitoring)"],
        ["/api/backup", "POST", "Excel/JSON backup export, Tally XML export"],
        ["/api/backup/emergency", "POST", "Emergency full backup"],
        ["/api/backup/restore", "POST", "Restore from backup file"],
        ["/api/auto-backup", "POST", "Auto-backup trigger after transactions"],
        ["/api/repair", "POST", "Data repair (merge wallets, fix owner roles)"],
        ["/api/tenants", "POST", "Tenant management (super-admin)"],
        ["/api/db-admin", "POST", "Database admin operations (super-admin)"],
        ["/api/admin/delete-account", "POST", "Account deletion (GDPR)"],
        ["/api/debug", "POST", "Debug endpoint (disabled in prod)"],
        ["/api/debug-env", "POST", "Env var debug (super-admin only)"],
        ["/api/debug-smtp", "POST", "SMTP connection test"],
        ["/api/einvoice", "POST", "E-invoice IRN generation via IRP"],
        ["/api/desktop-download", "POST", "Desktop installer download URLs"],
        ["/api/invoice-file", "POST", "Save invoice file reference"],
        ["/api/save-invoice-file", "POST", "Save invoice file content"],
        ["/api/cron/health-monitor", "GET", "Cron: health monitoring"],
        ["/api/cron/imap-scan", "GET", "Cron: IMAP payment screenshot scan"],
        ["/api/cron/sms-webhook", "POST", "SMS webhook receiver (UPI verification)"],
    ],
    col_widths=[4, 2.5, 8.5]
)
add_page_break()

# ============================================================
# 10. SECURITY ARCHITECTURE & INCIDENTS
# ============================================================
add_heading("10. Security Architecture & Incidents", level=1)

add_heading("10.1 Authentication & Authorization", level=2)
add_bullet("Session: JWT signed with SESSION_SECRET, stored in httpOnly cookie 'bizbook_session' (path=/, sameSite=lax)")
add_bullet("OTP: 6-digit, 10-minute expiry, single-use. Delivered via Brevo email (primary) + 2Factor SMS (fallback)")
add_bullet("Password hashing: bcrypt (cost factor 12)")
add_bullet("AFK auto-logout: 5 minutes of inactivity (extends if dialog is open)")
add_bullet("Remember Me: stores email in localStorage (not password)")

add_heading("10.2 Tenant Isolation", level=2)
add_bullet("Every table has tenantId column")
add_bullet("Every API route calls requireAuthAndTenant() or requireAuthAndRole() — access.tenantId from JWT, never body")
add_bullet("Startup safeguard: queries all tenants, asserts all ACTIVE, blocks startup if any are soft-deleted")
add_bullet("v4.114: hardcoded tenant list removed — ALL registered tenants are protected dynamically")

add_heading("10.3 Audit Trail", level=2)
add_bullet("Every create/update/delete writes an AuditLog entry: tenantId, userId, userName, action, entityType, entityId, entityName, changes (JSON diff), createdAt")
add_bullet("Audit log is append-only — no update or delete on AuditLog records")
add_bullet("Soft-delete: all data tables use isDeleted + deletedAt — nothing is truly deleted (reversible for 30 days)")

add_heading("10.4 Payment Security", level=2)
add_bullet("Razorpay with HMAC-SHA256 signature verification on every payment")
add_bullet("No card/UPI data touches BizBook Pro servers — Razorpay handles all sensitive data")
add_bullet("UPI verification: 3-layer (SMS webhook + IMAP scan + manual screenshot review)")

add_heading("10.5 Security Incidents & Remediation", level=2, color=(220, 38, 38))
add_para("INCIDENT 1: Brevo API key leaked in git history (commit a094565, 2026-06-20)", bold=True, color=(220, 38, 38))
add_bullet("Severity: HIGH — repo is public")
add_bullet("Key value exposed: [REDACTED — see Git history commit a094565 for full Brevo key]")
add_bullet("Status: Key is in git history (cannot be rewritten without force-push coordination). Rotation required at Brevo dashboard.")
add_bullet("Note: Brevo SSL issue previously blocked rotation — verify resolved before rotating.")

add_para("INCIDENT 2: Razorpay key secret — possible exposure in early commits", bold=True, color=(220, 38, 38))
add_bullet("Severity: MEDIUM — verify with: git log --all -p | grep -iE 'rzt_|razorpay.*secret'")
add_bullet("Status: Rotation recommended regardless of finding.")

add_para("INCIDENT 3: Personal email in user-facing text (v4.64.3)", bold=True, color=(220, 38, 38))
add_bullet("Severity: LOW — pranjalgoswamighy86@gmail.com was shown in sidebar/settings")
add_bullet("Status: FIXED in v4.64.3 — all personal info removed from user-facing text")

add_para("INCIDENT 4: Stale Service Worker cache (v4.155 stuck, users on old JS)", bold=True, color=(220, 38, 38))
add_bullet("Severity: MEDIUM — users saw old invoice CSS despite new deploys")
add_bullet("Status: FIXED in v4.178 — SW cache version bumped, forces browser refetch")

add_para("INCIDENT 5: Electron menu bar not updating after web deploys (v2.3.0)", bold=True, color=(220, 38, 38))
add_bullet("Severity: HIGH — 2-3 web deploys didn't reach desktop users")
add_bullet("Root cause: Electron cached old web app JS; menu-action listener was auth-gated")
add_bullet("Status: FIXED in v2.3.0 — desktop clears cache on every launch + global MenuActionBridge + VersionBadge for verification")
add_page_break()

# ============================================================
# 11. RBAC
# ============================================================
add_heading("11. Role-Based Access Control", level=1)
add_para("5-tier hierarchy. 4 tenant roles + 1 backend-only platform role. Enforced at API layer via requireAuthAndRole(roles[]).")

add_table(
    ["Role", "Scope", "Permissions", "Can Access Tenant Data?"],
    [
        ["VIEW_ONLY", "Tenant", "Read-only access to all modules. No create, edit, delete.", "Yes (read)"],
        ["DATA_ENTRY", "Tenant", "Create sales, purchases, expenses. Cannot delete or edit confirmed records. Cannot manage users.", "Yes"],
        ["JUNIOR_ADMIN", "Tenant", "Edit and correct entries. Full access to financial reports. Cannot manage subscription or users.", "Yes"],
        ["MAIN_ADMIN", "Tenant", "Owner of the tenant. Manages users, subscription, settings, backup. Full access.", "Yes"],
        ["SUPER_ADMIN", "Platform (backend only)", "Manages the SaaS itself — all tenants, super-admin panel, payment proofs. NOT available to tenants.", "All tenants (platform ops)"],
    ],
    col_widths=[3, 3, 7, 3]
)
add_para("Helper functions in src/store/app-store.ts:", bold=True)
add_code_block("canEdit(role)     → DATA_ENTRY, JUNIOR_ADMIN, MAIN_ADMIN\ncanCorrect(role)  → JUNIOR_ADMIN, MAIN_ADMIN\ncanManage(role)   → MAIN_ADMIN")
add_para("Known issue (v6.18.0): The sale-register Edit/Delete buttons use canCorrect() which excludes DATA_ENTRY. Per the product spec, DATA_ENTRY should be able to edit Quotation entries (but not delete them, and not edit confirmed invoices). This is a known gap — see Known Issues §13.", italic=True, color=(220, 38, 38))
add_page_break()

# ============================================================
# 12. SUBSCRIPTION & BILLING LOGIC
# ============================================================
add_heading("12. Subscription & Billing Logic", level=1)
add_para("Source of truth: src/app/api/subscription/route.ts. All pricing is hardcoded in the PLANS constant — update there and redeploy to change prices.")

add_heading("12.1 Free Tier", level=2)
add_bullet("One-time allocation on registration — NOT monthly renewal")
add_bullet("First 500 tenants: 100 free hours")
add_bullet("Tenants 501-10,000: 50 free hours")
add_bullet("After 10,000: 20 free hours")
add_bullet("Logic: getFreeTierHours(userCount) in subscription/route.ts")

add_heading("12.2 Paid Plans (5 tiers)", level=2)
add_table(
    ["Plan", "Hours", "MRP (₹)", "Discount %", "Customer Pays (₹)", "Role Allocation (Main/Junior/DE)"],
    [
        ["50Hrs Plan", "50", "749", "80%", "150", "10h / 15h / 25h"],
        ["100Hrs Plan", "100", "1,449", "85%", "217", "20h / 30h / 50h"],
        ["200Hrs Plan", "200", "2,849", "90%", "285", "40h / 60h / 100h"],
        ["500Hrs Plan", "500", "7,049", "93%", "493", "80h / 120h / 200h"],
        ["1000Hrs Plan", "1000", "14,049", "96%", "562", "40h / 60h / 100h"],
    ],
    col_widths=[3, 1.5, 2, 2, 2.5, 4]
)
add_para("NOTE: The 1000Hrs Plan has the same role allocation as the 200Hrs plan (40/60/100) — this may be a typo in the source. Verify with product owner before changing.", italic=True, color=(220, 38, 38))

add_heading("12.3 User Slots & Surcharges", level=2)
add_bullet("Default: 3 non-view-only users included with every plan")
add_bullet("Extra non-view-only user: ₹149 one-time fee (EXTRA_USER_ONE_TIME_FEE)")
add_bullet("Recharge surcharge: 15% added to plan price if tenant has extra non-view-only users (EXTRA_USER_RECHARGE_SURCHARGE_PERCENT)")
add_bullet("View-Only users: UNLIMITED and free (not counted toward any limit)")

add_heading("12.4 Wallet Model (v6.5+)", level=2)
add_bullet("Tenant-level wallet: one wallet, all companies under the tenant share the same hour balance")
add_bullet("Single main admin per tenant: cannot be replicated or transferred")
add_bullet("Global user count: uses distinct userId (main admin counted once)")
add_bullet("Payment enforcement: free tier allows 3 users; 4th user requires ₹149")
add_bullet("Usage limits: main admin can set per-user monthly hour limits")
add_bullet("Per-user hour tracking: UsageLog table records seconds used")
add_bullet("Auto-repair on startup: consolidates duplicate subscriptions, fixes broken owner roles, merges wallets")
add_page_break()

# ============================================================
# 13. KNOWN ISSUES & TECHNICAL DEBT
# ============================================================
add_heading("13. Known Issues & Technical Debt", level=1, color=(220, 38, 38))

add_heading("13.1 Quotation Feature — Data Integrity Gaps", level=2)
add_para("The quotation feature (v4.106) was implemented at the UI level but the financial isolation was never completed. Current behavior:", bold=True)
add_bullet("CREATE: A quotation (invoiceStatus=QUOTATION) STILL deducts inventory, creates debtors, posts journal entries, and creates receipts — same as a confirmed sale. There is NO skip-logic for quotations.")
add_bullet("CONFIRM-SALE: Only flips invoiceStatus to CONFIRMED. Does NOT re-trigger the side-effect pipeline (because nothing was skipped at create time).")
add_bullet("REPORTS: P&L, Balance Sheet, Day Report, Dashboard, GST Reports — NONE filter by invoiceStatus='CONFIRMED'. Quotation amounts ARE included in all financial reports.")
add_bullet("INVENTORY: Quotation quantities ARE deducted from inventory (because create always deducts).")
add_bullet("DEBTORS: Quotations with credit payment status DO create debtor records.")
add_bullet("RBAC: DATA_ENTRY users cannot edit quotations (UI uses canCorrect which excludes them). Per spec, they should be able to edit quotations but not delete them, and not edit confirmed invoices.")
add_para("Impact: Every quotation silently corrupts inventory counts and inflates financial reports. This is a CRITICAL data integrity issue.", bold=True, color=(220, 38, 38))
add_para("Fix plan (not yet implemented): See verification report in chat log around v6.18.0. Requires 5 coordinated changes: skip side-effects in create for quotations, trigger side-effects in confirm-sale, block DATA_ENTRY edits of confirmed sales, update reports to filter by invoiceStatus, fix frontend role gating.")

add_heading("13.2 TypeScript Errors (pre-existing, non-blocking)", level=2)
add_bullet("src/app/app/page.tsx: 'tds-register' and 'payment-proof-review' not in ViewType union (line 87, 92)")
add_bullet("src/app/api/ai-smart-search/route.ts: Request vs NextRequest type mismatch")
add_bullet("src/app/api/autocomplete/route.ts: autocompleteIndex property missing on Prisma client")
add_bullet("src/components/modules/company-select.tsx: BackupImportDialog missing onOpenChange prop")
add_para("These are all pre-existing and do not block the build (Next.js compiles despite them). Fix when convenient.")

add_heading("13.3 Service Worker Cache Versioning", level=2)
add_bullet("SW cache version must be manually bumped in src/lib/version.ts on every deploy that changes invoice CSS or print HTML")
add_bullet("If forgotten, users see stale cached JS — this happened at v4.155 (stuck for 5 days until v4.178)")
add_bullet("Consider automating: bump SW version on every build via postbuild.js")

add_heading("13.4 PM2 Cluster Mode", level=2)
add_bullet("PM2 sometimes fails to start in Railway container — falls back to direct server.js (see startup log)")
add_bullet("Not critical (single instance works) but loses the zero-downtime benefit")
add_bullet("Investigate: PM2 may need --no-daemon flag or different config in Docker")

add_heading("13.5 Brevo SSL Issue", level=2)
add_bullet("Brevo API was returning SSL errors at some point — blocked key rotation")
add_bullet("Verify resolved before attempting rotation")
add_bullet("If still broken, use Resend as primary email provider instead")

add_heading("13.6 Desktop App — Static File Bundling", level=2)
add_bullet("Electron app loads Railway URL directly (v5.6+) — no local static files needed")
add_bullet("But: first load requires internet. If offline, app shows blank window with retry logic (10 attempts, 1.5s apart)")
add_bullet("Consider: bundle a minimal offline fallback page in the Electron app")
add_page_break()

# ============================================================
# 14. CRITICAL FILES REFERENCE
# ============================================================
add_heading("14. Critical Files Reference", level=1)
add_para("The following files are the single source of truth for their respective concerns. Change them carefully.")

add_table(
    ["File", "Purpose", "Change Frequency"],
    [
        ["src/lib/version.ts", "APP_VERSION + APP_BUILD_DATE (single source of truth)", "Every release"],
        ["prisma/schema.prisma", "Database schema (30+ models)", "Rare — requires db push"],
        ["src/app/api/subscription/route.ts", "PLANS array, free tier logic, wallet model, pricing", "When pricing changes"],
        ["src/app/api/sales/route.ts", "Sale CRUD, confirm-sale, inventory deduction, JE creation", "Feature work"],
        ["src/app/api/auth/route.ts", "Login, register, OTP, password reset, user creation", "Security work"],
        ["src/app/api/repair/route.ts", "Auto-repair: merge wallets, fix owner roles", "Startup + manual"],
        ["src/lib/multi-ai.ts", "4-provider AI fallback chain", "Rare"],
        ["src/lib/db-soft-delete.ts", "Prisma client with soft-delete extension", "Rare"],
        ["src/lib/api-helpers.ts", "requireAuthAndTenant, requireAuthAndRole, writeAuditLog", "Rare"],
        ["src/lib/auth-fetch.ts", "Authenticated fetch wrapper (sends cookie)", "Rare"],
        ["src/store/app-store.ts", "Zustand store: user, tenant, currentView, canEdit/canCorrect/canManage", "Feature work"],
        ["src/app/app/page.tsx", "Main app page: module router, keyboard shortcuts, session validation", "Feature work"],
        ["src/app/layout.tsx", "Root layout: MenuActionBridge, VersionBadge, Toaster, OfflineBanner", "Rare"],
        ["src/components/app/menu-action-bridge.tsx", "Global Electron menu-action handler (always mounted)", "Rare"],
        ["src/components/app/version-badge.tsx", "Visible version badge + Electron ping", "Rare"],
        ["src/components/app/help-modal.tsx", "Help & Support modal (FAQ, Guides, AI Chat, Manage tabs)", "Content updates"],
        ["src/components/app/help-chat.tsx", "AI Support Chat tab component", "Rare"],
        ["src/lib/help-chat-trigger.ts", "Global openHelpChat() event dispatcher (F1)", "Rare"],
        ["src/components/app/sidebar.tsx", "Sidebar nav, help button, backup dialog", "Feature work"],
        ["src/components/modules/sale-register.tsx", "Sale register UI (largest component, ~1500 lines)", "Feature work"],
        ["electron/main.ts", "Electron main process: window, menu, IPC, printing, cache clear", "Desktop releases"],
        ["electron/preload.ts", "Electron preload: contextBridge API surface", "Desktop releases"],
        ["electron-builder.yml", "Desktop build config (NSIS, DMG, AppImage targets)", "Rare"],
        [".github/workflows/build-desktop.yml", "GitHub Actions: build desktop on tag push", "Rare"],
        ["scripts/railway-start.js", "Railway startup: prisma generate, db push, tenant protect, PM2 start", "Rare"],
        ["ecosystem.config.js", "PM2 cluster config (2 instances)", "Rare"],
        ["Dockerfile", "Railway Docker build (9 steps)", "Rare"],
        ["postbuild.js", "Post-build: copy static files to standalone, stamp version", "Rare"],
    ],
    col_widths=[6, 7, 3]
)
add_page_break()

# ============================================================
# 15. OPERATIONAL RUNBOOK
# ============================================================
add_heading("15. Operational Runbook", level=1)

add_heading("15.1 Deploy a Web App Update", level=2)
add_bullet("1. Make code changes locally")
add_bullet("2. Bump version in src/lib/version.ts (APP_VERSION + APP_BUILD_DATE)")
add_bullet("3. If invoice CSS or print HTML changed: bump SW cache version too")
add_bullet("4. git add -A && git commit -m 'vX.Y.Z: description'")
add_bullet("5. git push origin main")
add_bullet("6. Railway auto-rebuilds (~3-5 min). Monitor at Railway dashboard")
add_bullet("7. Verify: curl -s https://carefree-success-production-7766.up.railway.app/ | grep -oE 'v[0-9]+\\.[0-9]+\\.[0-9]+'")
add_bullet("8. Desktop app picks up changes on next launch (cache cleared automatically since v2.3.0)")

add_heading("15.2 Deploy a Desktop App Update", level=2)
add_bullet("1. Web app changes must be deployed and verified first (see 15.1)")
add_bullet("2. If Electron code (electron/main.ts, electron/preload.ts) changed: bump version in package.json")
add_bullet("3. git tag vX.Y.Z -m 'Desktop vX.Y.Z: description' && git push origin vX.Y.Z")
add_bullet("4. GitHub Actions triggers automatically (tag v* pattern)")
add_bullet("5. Build takes ~10-15 min (Windows + macOS + Linux matrix)")
add_bullet("6. Release appears as DRAFT at https://github.com/pranjalgoswamighy86/bizbook-pro/releases")
add_bullet("7. Review the draft, then publish to make installers public")
add_bullet("8. Users download new installer; old version auto-updates if electron-updater is configured")

add_heading("15.3 Emergency: Web App Down", level=2)
add_bullet("1. Check Railway dashboard → carefree-success → Deployments (latest status)")
add_bullet("2. Check logs: Railway → carefree-success → Logs")
add_bullet("3. Common causes: build failure (check build log), OOM (increase NODE_OPTIONS), Prisma error (check DATABASE_URL)")
add_bullet("4. Quick rollback: Railway → Deployments → click previous successful deploy → Redeploy")
add_bullet("5. If DB issue: check /api/health endpoint, verify DATABASE_URL env var, run prisma db push")
add_bullet("6. If all else fails: the desktop app still works (it loads the Railway URL — if Railway is down, desktop shows retry screen)")

add_heading("15.4 Emergency: Data Corruption", level=2)
add_bullet("1. Startup backup: every startup creates /tmp/bizbook-backups/bizbook_startup_<timestamp>.json (1.8MB, 2631 records as of last log)")
add_bullet("2. Auto-backup: every transaction triggers an Excel backup download client-side")
add_bullet("3. Manual backup: Settings → Backup & Restore → Export (JSON or Tally XML)")
add_bullet("4. Repair endpoint: POST /api/repair with action='full-repair' (merges wallets, fixes owner roles)")
add_bullet("5. Last resort: restore from startup backup JSON via /api/backup/restore")

add_heading("15.5 Add a New API Endpoint", level=2)
add_bullet("1. Create src/app/api/<name>/route.ts")
add_bullet("2. Import requireAuthAndTenant, requireAuthAndRole, writeAuditLog from '@/lib/api-helpers'")
add_bullet("3. Import db from '@/lib/db-soft-delete' (NOT '@/lib/db' — soft-delete extension required)")
add_bullet("4. At top of each action: const access = await requireAuthAndTenant(req, tenantId); if (access instanceof NextResponse) return access;")
add_bullet("5. Use access.tenantId (NOT body.tenantId) for all DB queries")
add_bullet("6. Wrap multi-step writes in db.$transaction(async (tx) => { ... })")
add_bullet("7. Write audit log inside the transaction")
add_bullet("8. Handle errors: return NextResponse.json({error: msg}, {status: 4xx/5xx})")
add_bullet("9. Test with curl before deploying")

add_heading("15.6 Rotate a Compromised API Key", level=2)
add_bullet("1. Visit the provider's dashboard (Brevo/Razorpay/etc.)")
add_bullet("2. Delete the old key (or revoke it)")
add_bullet("3. Create a new key")
add_bullet("4. Railway → carefree-success → Variables → update the env var")
add_bullet("5. Railway → Deployments → Redeploy (so new env var takes effect)")
add_bullet("6. Verify: check /api/health and test the affected feature")
add_bullet("7. NOTE: Git history still contains the old key. Rotation is the only remediation — do not attempt to rewrite history.")
add_page_break()

# ============================================================
# 16. SECURITY ROTATION CHECKLIST
# ============================================================
add_heading("16. Security Rotation Checklist", level=1, color=(220, 38, 38))
add_para("The following credentials have been exposed or are due for rotation. Complete each item and check it off.", bold=True, color=(220, 38, 38))

add_table(
    ["Credential", "Status", "Action Required", "Priority"],
    [
        ["BREVO_API_KEY", "LEAKED in git history (commit a094565)", "Rotate at Brevo dashboard → update Railway env var → redeploy", "CRITICAL"],
        ["RAZORPAY_KEY_SECRET", "Possibly exposed in early commits", "Verify with git log -p | grep razorpay; rotate regardless", "HIGH"],
        ["SESSION_SECRET", "Unknown exposure", "Rotate: generate new with openssl rand -hex 32; update Railway; redeploy (logs out all users)", "MEDIUM"],
        ["GEMINI_API_KEY", "Not exposed but good practice", "Rotate quarterly", "LOW"],
        ["OPENAI_API_KEY", "Not exposed but good practice", "Rotate quarterly", "LOW"],
        ["SMTP_PASS (Gmail App Password)", "Not exposed", "Rotate if Gmail account compromised", "LOW"],
        ["ADMIN_EMAIL password", "Test password '789456' was mentioned in chat (v1.x)", "CHANGE IMMEDIATELY — login to admin@bizbook.pro with old password, change in Settings", "CRITICAL"],
    ],
    col_widths=[4, 4, 6, 2]
)

add_hr()
add_para("END OF DOCUMENT", bold=True, size=12, color=(16, 110, 88))
add_para("This document is classified INTERNAL — BACKEND TEAM ONLY. Do not distribute externally. Do not commit to public repositories. If you find this document outside the Tahigo International backend team, notify security@tahigo.in immediately.", italic=True, size=10, color=(220, 38, 38))

# ---- Save ----
doc.save(OUTPUT)
print(f"✅ Saved: {OUTPUT}")
print(f"   Size: {os.path.getsize(OUTPUT) / 1024:.1f} KB")
